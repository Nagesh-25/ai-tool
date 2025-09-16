"""
Document processing service for extracting text from various file formats
"""

import io
import logging
from typing import Optional, Dict, Any
import asyncio
from concurrent.futures import ThreadPoolExecutor

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import cv2
import numpy as np

# Google Cloud Vision for OCR
from google.cloud import vision
from google.cloud import documentai

from app.models.schemas import DocumentType
from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document formats and extracting text."""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=4)
        self.vision_client = None
        self.documentai_client = None
        
        # Initialize Google Cloud Vision if credentials are available
        try:
            self.vision_client = vision.ImageAnnotatorClient()
        except Exception as e:
            logger.warning(f"Google Cloud Vision not available: {e}")
        
        # Initialize Document AI client if credentials are available
        try:
            self.documentai_client = documentai.DocumentProcessorServiceClient()
        except Exception as e:
            logger.warning(f"Google Cloud Document AI not available: {e}")
    
    async def extract_text(
        self, 
        document_id: str, 
        file_type: DocumentType, 
        storage_path: str
    ) -> Optional[str]:
        """
        Extract text from a document based on its type.
        
        Args:
            document_id: Unique document identifier
            file_type: Type of document (PDF, DOC, DOCX, IMAGE, TEXT)
            storage_path: Path to the document in cloud storage
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            # Download document content from storage
            document_content = await self._download_document(storage_path)
            if not document_content:
                logger.error(f"Failed to download document {document_id}")
                return None
            
            # Extract text based on document type
            if file_type == DocumentType.PDF:
                return await self._extract_pdf_text(document_content)
            elif file_type in [DocumentType.DOC, DocumentType.DOCX]:
                return await self._extract_docx_text(document_content)
            elif file_type == DocumentType.IMAGE:
                return await self._extract_image_text(document_content)
            elif file_type == DocumentType.TEXT:
                return document_content.decode('utf-8', errors='ignore')
            else:
                logger.error(f"Unsupported document type: {file_type}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from document {document_id}: {e}")
            return None
    
    async def _download_document(self, storage_path: str) -> Optional[bytes]:
        """Download document content from cloud storage."""
        try:
            from app.services.storage_service import StorageService
            storage_service = StorageService()
            await storage_service.initialize()
            return await storage_service.download_document(storage_path)
        except Exception as e:
            logger.error(f"Error downloading document from {storage_path}: {e}")
            return None
    
    async def _extract_pdf_text(self, pdf_content: bytes) -> Optional[str]:
        """Extract text from PDF using multiple methods for better accuracy."""
        try:
            # Method 1: Try Google Cloud Document AI first (best for complex documents)
            if self.documentai_client:
                text = await self._extract_pdf_with_documentai(pdf_content)
                if text and len(text.strip()) > 100:
                    return text
            
            # Method 2: Try pdfplumber (better for complex layouts)
            text = await self._extract_pdf_with_pdfplumber(pdf_content)
            if text and len(text.strip()) > 100:  # Minimum text threshold
                return text
            
            # Method 3: Fallback to PyPDF2
            text = await self._extract_pdf_with_pypdf2(pdf_content)
            if text and len(text.strip()) > 100:
                return text
            
            # Method 4: If all fail, try OCR on PDF pages
            return await self._extract_pdf_with_ocr(pdf_content)
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    async def _extract_pdf_with_pdfplumber(self, pdf_content: bytes) -> Optional[str]:
        """Extract text using pdfplumber."""
        def _extract():
            try:
                with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    return '\n\n'.join(text_parts)
            except Exception as e:
                logger.error(f"pdfplumber extraction failed: {e}")
                return None
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    async def _extract_pdf_with_pypdf2(self, pdf_content: bytes) -> Optional[str]:
        """Extract text using PyPDF2."""
        def _extract():
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
                text_parts = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return '\n\n'.join(text_parts)
            except Exception as e:
                logger.error(f"PyPDF2 extraction failed: {e}")
                return None
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    async def _extract_pdf_with_documentai(self, pdf_content: bytes) -> Optional[str]:
        """Extract text using Google Cloud Document AI."""
        def _extract():
            try:
                # Note: This requires a Document AI processor to be set up
                # For now, we'll use a simplified approach
                # In production, you would configure a specific processor
                logger.info("Document AI extraction not fully configured - using fallback methods")
                return None
            except Exception as e:
                logger.error(f"Document AI extraction failed: {e}")
                return None
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    async def _extract_pdf_with_ocr(self, pdf_content: bytes) -> Optional[str]:
        """Extract text from PDF using OCR (for scanned PDFs)."""
        try:
            # Convert PDF pages to images and then OCR
            # This is a simplified version - in production, you'd use pdf2image
            logger.info("Attempting OCR extraction for PDF")
            # TODO: Implement PDF to image conversion and OCR
            return None
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return None
    
    async def _extract_docx_text(self, docx_content: bytes) -> Optional[str]:
        """Extract text from DOC/DOCX files."""
        def _extract():
            try:
                doc = Document(io.BytesIO(docx_content))
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(' | '.join(row_text))
                
                return '\n\n'.join(text_parts)
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                return None
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    async def _extract_image_text(self, image_content: bytes) -> Optional[str]:
        """Extract text from images using OCR."""
        try:
            # Method 1: Try Google Cloud Vision API first (best accuracy)
            if self.vision_client:
                text = await self._extract_with_google_vision(image_content)
                if text and len(text.strip()) > 10:
                    logger.info("Successfully extracted text using Google Cloud Vision")
                    return text
            
            # Method 2: Fallback to Tesseract OCR
            logger.info("Falling back to Tesseract OCR")
            return await self._extract_with_tesseract(image_content)
            
        except Exception as e:
            logger.error(f"Image OCR extraction failed: {e}")
            return None
    
    async def _extract_with_google_vision(self, image_content: bytes) -> Optional[str]:
        """Extract text using Google Cloud Vision API."""
        def _extract():
            try:
                image = vision.Image(content=image_content)
                response = self.vision_client.text_detection(image=image)
                texts = response.text_annotations
                
                if texts:
                    # Return the first (full) text annotation
                    return texts[0].description
                return None
            except Exception as e:
                logger.error(f"Google Vision OCR failed: {e}")
                return None
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    async def _extract_with_tesseract(self, image_content: bytes) -> Optional[str]:
        """Extract text using Tesseract OCR."""
        def _extract():
            try:
                # Load and preprocess image
                image = Image.open(io.BytesIO(image_content))
                
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Convert to numpy array for OpenCV processing
                img_array = np.array(image)
                
                # Preprocess image for better OCR
                processed_img = self._preprocess_image_for_ocr(img_array)
                
                # Extract text using Tesseract
                text = pytesseract.image_to_string(
                    processed_img,
                    config='--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,!?;:()[]{}"\' '
                )
                
                return text.strip()
            except Exception as e:
                logger.error(f"Tesseract OCR failed: {e}")
                return None
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    def _preprocess_image_for_ocr(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image to improve OCR accuracy."""
        try:
            # Convert to grayscale
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
            else:
                gray = image
            
            # Apply Gaussian blur to reduce noise
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            
            # Apply adaptive thresholding
            thresh = cv2.adaptiveThreshold(
                blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Morphological operations to clean up the image
            kernel = np.ones((1, 1), np.uint8)
            cleaned = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            return cleaned
        except Exception as e:
            logger.error(f"Image preprocessing failed: {e}")
            return image
    
    async def get_document_metadata(self, document_content: bytes, file_type: DocumentType) -> Dict[str, Any]:
        """Extract metadata from document."""
        metadata = {
            "file_type": file_type.value,
            "size_bytes": len(document_content),
            "extraction_method": None,
            "language_detected": None,
            "page_count": None,
            "has_images": False,
            "has_tables": False,
            "ocr_confidence": None,
            "processing_notes": []
        }
        
        try:
            if file_type == DocumentType.PDF:
                metadata.update(await self._get_pdf_metadata(document_content))
            elif file_type in [DocumentType.DOC, DocumentType.DOCX]:
                metadata.update(await self._get_docx_metadata(document_content))
            elif file_type == DocumentType.IMAGE:
                metadata.update(await self._get_image_metadata(document_content))
            
            # Add processing notes
            if metadata.get("extraction_method") == "ocr":
                metadata["processing_notes"].append("Document processed using OCR")
            elif metadata.get("extraction_method") == "google_vision":
                metadata["processing_notes"].append("Document processed using Google Cloud Vision")
            elif metadata.get("extraction_method") == "documentai":
                metadata["processing_notes"].append("Document processed using Google Cloud Document AI")
                
        except Exception as e:
            logger.error(f"Error extracting document metadata: {e}")
            metadata["processing_notes"].append(f"Metadata extraction error: {str(e)}")
        
        return metadata
    
    async def _get_pdf_metadata(self, pdf_content: bytes) -> Dict[str, Any]:
        """Extract PDF-specific metadata."""
        def _extract():
            try:
                with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                    return {
                        "page_count": len(pdf.pages),
                        "has_images": any(len(page.images) > 0 for page in pdf.pages),
                        "has_tables": any(len(page.find_tables()) > 0 for page in pdf.pages),
                        "extraction_method": "pdfplumber"
                    }
            except Exception as e:
                logger.error(f"PDF metadata extraction failed: {e}")
                return {}
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    async def _get_docx_metadata(self, docx_content: bytes) -> Dict[str, Any]:
        """Extract DOCX-specific metadata."""
        def _extract():
            try:
                doc = Document(io.BytesIO(docx_content))
                return {
                    "page_count": None,  # DOCX doesn't have fixed page count
                    "has_images": any(len(paragraph._element.xpath('.//a:blip')) > 0 for paragraph in doc.paragraphs),
                    "has_tables": len(doc.tables) > 0,
                    "extraction_method": "python-docx"
                }
            except Exception as e:
                logger.error(f"DOCX metadata extraction failed: {e}")
                return {}
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    async def _get_image_metadata(self, image_content: bytes) -> Dict[str, Any]:
        """Extract image-specific metadata."""
        def _extract():
            try:
                image = Image.open(io.BytesIO(image_content))
                metadata = {
                    "page_count": 1,
                    "has_images": True,
                    "has_tables": False,
                    "extraction_method": "ocr",
                    "image_dimensions": image.size,
                    "image_format": image.format,
                    "ocr_confidence": None
                }
                
                # Try to get OCR confidence from Google Vision if available
                if self.vision_client:
                    try:
                        vision_image = vision.Image(content=image_content)
                        response = self.vision_client.text_detection(image=vision_image)
                        if response.text_annotations:
                            # Calculate average confidence from text annotations
                            confidences = []
                            for annotation in response.text_annotations[1:]:  # Skip first (full text)
                                if hasattr(annotation, 'confidence'):
                                    confidences.append(annotation.confidence)
                            if confidences:
                                metadata["ocr_confidence"] = sum(confidences) / len(confidences)
                                metadata["extraction_method"] = "google_vision"
                    except Exception as e:
                        logger.warning(f"Could not get OCR confidence from Vision API: {e}")
                
                return metadata
            except Exception as e:
                logger.error(f"Image metadata extraction failed: {e}")
                return {}
        
        return await asyncio.get_event_loop().run_in_executor(self.executor, _extract)
    
    def cleanup(self):
        """Cleanup resources."""
        if self.executor:
            self.executor.shutdown(wait=True)
